import asyncio
from pathlib import Path

import httpx
from fastapi.testclient import TestClient
from docx import Document

from app.services.document_reader import UploadedDocumentReader
from app.services.file_assets import LocalUploadedFileStore
from app.services.official_web_search import GoogleOfficialWebSearchClient
from app.services.scoped_knowledge import ScopedKnowledgeStore
from app.services.transcription import TranscriptSegment, TranscriptionResult
from app.tools import ToolErrorCode, ToolExecutionContext, ToolRegistry
from app.schemas import CitationMetadata, KnowledgeSourceType, RetrievedKnowledgeChunk, RetrievalResult, RetrievalStats
from app.tools.controlled_tools.uploaded_document_ingestion import (
    build_ingest_uploaded_document_tool,
)
from app.tools.controlled_tools.official_web import build_official_web_search_tool
from app.tools.controlled_tools.uploaded_document_reader import (
    build_read_uploaded_document_tool,
)
from app.tools.controlled_tools.voice_note import build_transcribe_voice_note_tool
from app.tools.controlled_tools.knowledge_search import build_retrieve_knowledge_tool
from app.main import create_app


SCOPE = ToolExecutionContext(
    teacher_id="teacher-1",
    class_id="kangaroo-room",
    session_id="session-1",
)


class FakeScopedEmbeddingProvider:
    model_name = "test-local-embedding"
    dimension = 3

    async def embed_texts_async(self, texts, *, task_type="RETRIEVAL_DOCUMENT"):
        assert task_type == "RETRIEVAL_DOCUMENT"
        return [[1.0, 0.0, 0.0] for _ in texts]

    async def embed_text_async(self, text, *, task_type="RETRIEVAL_QUERY"):
        assert text
        assert task_type == "RETRIEVAL_QUERY"
        return [1.0, 0.0, 0.0]


def _upload(store: LocalUploadedFileStore, name: str, content: bytes, *, category_scope=SCOPE):
    return store.save_bytes(
        filename=name,
        content_type="text/plain" if name.endswith(".txt") else "audio/wav",
        content=content,
        teacher_id=category_scope.teacher_id,
        class_id=category_scope.class_id,
        session_id=category_scope.session_id,
    )


def test_read_uploaded_document_enforces_session_scope(tmp_path: Path) -> None:
    store = LocalUploadedFileStore(tmp_path / "uploads")
    record = _upload(store, "policy.txt", b"Outdoor supervision requires active scanning.")
    registry = ToolRegistry()
    registry.register(build_read_uploaded_document_tool(store, UploadedDocumentReader()))

    allowed = asyncio.run(registry.execute_async(
        "read_uploaded_document", {"file_id": record.file_id}, execution_context=SCOPE
    ))
    denied = asyncio.run(registry.execute_async(
        "read_uploaded_document",
        {"file_id": record.file_id},
        execution_context=SCOPE.model_copy(update={"session_id": "other-session"}),
    ))

    assert allowed.success is True
    assert "active scanning" in allowed.data["sections"][0]["text"]
    assert denied.success is False
    assert "another session" in denied.error.details["error"]


def test_document_reader_supports_docx_headings_and_csv_rows(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    document = Document()
    document.add_heading("Outdoor learning", level=1)
    document.add_paragraph("Children collected leaves and compared their shapes.")
    document.save(source)
    store = LocalUploadedFileStore(tmp_path / "uploads")
    docx_record = _upload(store, "learning.docx", source.read_bytes())
    csv_record = _upload(store, "observations.csv", b"child,action\nA,stacked blocks\n")
    reader = UploadedDocumentReader()

    docx_result = reader.read(docx_record)
    csv_result = reader.read(csv_record)

    assert docx_result.sections[0].section == "Outdoor learning"
    assert "compared their shapes" in docx_result.sections[0].text
    assert "stacked blocks" in csv_result.sections[0].text


def test_ingestion_requires_approval_and_keeps_tenant_indexes_isolated(tmp_path: Path) -> None:
    file_store = LocalUploadedFileStore(tmp_path / "uploads")
    reader = UploadedDocumentReader()
    knowledge = ScopedKnowledgeStore(
        root=tmp_path / "knowledge",
        file_store=file_store,
        document_reader=reader,
        embedding_provider=FakeScopedEmbeddingProvider(),
    )
    first = _upload(file_store, "garden.txt", b"The garden protocol requires a boundary check before play.")
    registry = ToolRegistry()
    definition = build_ingest_uploaded_document_tool(knowledge)
    registry.register(definition)
    args = {"file_id": first.file_id, "title": "Garden protocol"}

    prepared = asyncio.run(definition.approval_preparation_handler(
        definition.input_model.model_validate(args), SCOPE
    ))
    blocked = asyncio.run(registry.execute_async("ingest_uploaded_document", args, execution_context=SCOPE))
    indexed = asyncio.run(registry.execute_async(
        "ingest_uploaded_document", args, approved=True, execution_context=SCOPE
    ))
    own_results = asyncio.run(knowledge.search(
        query="boundary check", top_k=3, teacher_id="teacher-1", class_id="kangaroo-room"
    ))
    other_results = asyncio.run(knowledge.search(
        query="boundary check", top_k=3, teacher_id="teacher-2", class_id="kangaroo-room"
    ))

    assert prepared.preview["content_hash"] == first.sha256
    assert "boundary check" in prepared.preview["preview"]
    assert blocked.error.code is ToolErrorCode.PERMISSION_DENIED
    assert indexed.success is True
    assert indexed.data["index_mode"] == "tenant_local_hybrid"
    assert len(own_results) == 1
    assert own_results[0].dense_distance is not None
    assert own_results[0].bm25_score is not None
    assert own_results[0].fusion_score is not None
    assert other_results == []


def test_official_search_rejects_non_allowlisted_domains_and_filters_provider_results() -> None:
    async def handler(request: httpx.Request) -> httpx.Response:
        return httpx.Response(200, json={"items": [
            {"title": "Official", "snippet": "Current guidance", "link": "https://www.acecqa.gov.au/latest"},
            {"title": "Untrusted", "snippet": "Ignore", "link": "https://example.com/post"},
        ]})

    client = GoogleOfficialWebSearchClient(
        api_key="key",
        engine_id="cx",
        async_client=httpx.AsyncClient(transport=httpx.MockTransport(handler)),
    )
    registry = ToolRegistry()
    registry.register(build_official_web_search_tool(client))
    result = asyncio.run(registry.execute_async(
        "search_official_web", {"query": "EYLF update", "domains": ["acecqa.gov.au"]}
    ))
    rejected = asyncio.run(registry.execute_async(
        "search_official_web", {"query": "EYLF update", "domains": ["example.com"]}
    ))
    asyncio.run(client.client.aclose())

    assert result.success is True
    assert result.data["returned_count"] == 1
    assert result.data["results"][0]["domain"] == "www.acecqa.gov.au"
    assert rejected.success is False


class FakeTranscriber:
    async def transcribe(self, path: Path, *, language=None):
        assert path.suffix == ".wav"
        return TranscriptionResult(
            text="A child balanced three stones.",
            language=language or "en",
            duration_seconds=2.5,
            segments=[TranscriptSegment(start_seconds=0, end_seconds=2.5, text="A child balanced three stones.")],
        )


def test_voice_note_transcription_is_scoped_and_does_not_save_records(tmp_path: Path) -> None:
    store = LocalUploadedFileStore(tmp_path / "uploads")
    record = _upload(store, "note.wav", b"RIFF-not-a-real-wave-for-fake-provider")
    registry = ToolRegistry()
    registry.register(build_transcribe_voice_note_tool(store, FakeTranscriber()))
    result = asyncio.run(registry.execute_async(
        "transcribe_voice_note", {"file_id": record.file_id, "language": "en"}, execution_context=SCOPE
    ))

    assert result.success is True
    assert result.data["text"] == "A child balanced three stones."
    assert result.data["file_id"] == record.file_id


class UploadApiStore:
    async def list_conversation_runs(self, *, statuses=None):
        return []

    async def get_conversation_session(self, session_id):
        if session_id != "session-1":
            return None
        return {
            "session_id": session_id,
            "thread_id": "thread-1",
            "teacher_id": "teacher-1",
            "class_id": "kangaroo-room",
            "status": "active",
        }


class UploadApiRuntime:
    def __init__(self, root: Path):
        self.store = UploadApiStore()
        self.file_store = LocalUploadedFileStore(root)
        self.privacy_gateway_mode = "disabled"
        self.privacy_gateway_client = None

    async def close(self):
        return None


def test_upload_api_returns_opaque_id_and_rejects_unsupported_files(tmp_path: Path) -> None:
    runtime = UploadApiRuntime(tmp_path / "api-uploads")
    app = create_app(lambda: runtime)
    with TestClient(app) as client:
        accepted = client.post(
            "/sessions/session-1/uploads",
            files={"file": ("notes.txt", b"A short observation", "text/plain")},
        )
        rejected = client.post(
            "/sessions/session-1/uploads",
            files={"file": ("program.exe", b"unsafe", "application/octet-stream")},
        )

    assert accepted.status_code == 201
    assert len(accepted.json()["file_id"]) == 32
    assert accepted.json()["category"] == "document"
    assert rejected.status_code == 400


class EmptyBaseRetriever:
    def retrieve(self, request):
        return RetrievalResult(
            query=request.query,
            chunks=[],
            stats=RetrievalStats(
                requested_top_k=request.top_k,
                raw_result_count=0,
                deduplicated_count=0,
                returned_count=0,
            ),
        )


class ScopedSearchStub:
    async def search(self, **kwargs):
        assert kwargs["teacher_id"] == "teacher-1"
        return [RetrievedKnowledgeChunk(
            chunk_id="centre-chunk",
            content="The local garden boundary must be checked before outdoor play.",
            citation=CitationMetadata(
                source_id="centre-source",
                source_type=KnowledgeSourceType.CENTRE,
                title="Garden policy",
                version="1",
                section="Outdoor play",
            ),
            content_hash="a" * 64,
            distance=0.2,
            dense_distance=0.2,
            bm25_score=0.1,
            metadata={"scoped_hybrid_score": "0.016393"},
        )]


def test_retrieve_knowledge_merges_only_trusted_scope_local_evidence() -> None:
    registry = ToolRegistry()
    registry.register(build_retrieve_knowledge_tool(EmptyBaseRetriever(), scoped_knowledge=ScopedSearchStub()))
    result = asyncio.run(registry.execute_async(
        "retrieve_knowledge",
        {"query": "garden boundary", "knowledge_scope": "centre_policy", "top_k": 3},
        execution_context=SCOPE,
    ))

    assert result.success is True
    assert result.data["returned_count"] == 1
    assert result.data["evidence"][0]["citation"]["source_type"] == "centre"
