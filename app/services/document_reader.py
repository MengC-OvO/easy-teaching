"""Deterministic parsing for files admitted by the scoped upload store."""

from __future__ import annotations

import csv
from pathlib import Path
from typing import List, Optional

from pydantic import BaseModel, Field

from app.services.file_assets import UploadedFileRecord
from app.services.knowledge_ingestion import ParsedTextBlock


class DocumentSection(BaseModel):
    text: str
    section: Optional[str] = None
    page: Optional[int] = Field(default=None, ge=1)


class DocumentReadResult(BaseModel):
    file_id: str
    filename: str
    content_type: str
    sections: List[DocumentSection]
    extracted_chars: int = Field(ge=0)
    truncated: bool = False


class UploadedDocumentReader:
    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".csv"}

    def read(
        self,
        record: UploadedFileRecord,
        *,
        max_chars: int = 20_000,
    ) -> DocumentReadResult:
        if record.category != "document":
            raise ValueError("Uploaded file is not a document")
        path = Path(record.stored_path)
        extension = path.suffix.casefold()
        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported document extension: {extension}")
        blocks = self.to_parsed_blocks(record)
        sections: List[DocumentSection] = []
        remaining = max_chars
        truncated = False
        for block in blocks:
            if remaining <= 0:
                truncated = True
                break
            text = block.content.strip()
            if len(text) > remaining:
                text = text[:remaining].rstrip()
                truncated = True
            if text:
                sections.append(
                    DocumentSection(
                        text=text,
                        section=block.section,
                        page=block.page,
                    )
                )
                remaining -= len(text)
            if truncated:
                break
        return DocumentReadResult(
            file_id=record.file_id,
            filename=record.original_name,
            content_type=record.content_type,
            sections=sections,
            extracted_chars=sum(len(item.text) for item in sections),
            truncated=truncated,
        )

    def to_parsed_blocks(self, record: UploadedFileRecord) -> List[ParsedTextBlock]:
        path = Path(record.stored_path)
        extension = path.suffix.casefold()
        if extension == ".pdf":
            return self._pdf_blocks(path)
        if extension == ".docx":
            return self._docx_blocks(path)
        if extension == ".csv":
            return self._csv_blocks(path)
        if extension in {".txt", ".md"}:
            text = path.read_text(encoding="utf-8-sig")
            return [
                ParsedTextBlock(
                    content=text,
                    section="Uploaded document",
                    metadata={"parser": "plain_text"},
                )
            ]
        raise ValueError(f"Unsupported document extension: {extension}")

    @staticmethod
    def _pdf_blocks(path: Path) -> List[ParsedTextBlock]:
        import pymupdf
        import pymupdf4llm

        with pymupdf.open(path) as document:
            pages = list(range(document.page_count))
        chunks = pymupdf4llm.to_markdown(
            str(path),
            pages=pages,
            page_chunks=True,
            use_ocr=False,
            force_ocr=False,
            force_text=True,
            header=False,
            footer=False,
            write_images=False,
            embed_images=False,
            show_progress=False,
        )
        blocks: List[ParsedTextBlock] = []
        for index, chunk in enumerate(chunks, start=1):
            text = str(chunk.get("text") or "").strip()
            if text:
                metadata = chunk.get("metadata") or {}
                blocks.append(
                    ParsedTextBlock(
                        content=text,
                        section="Uploaded PDF",
                        page=int(metadata.get("page_number") or index),
                        metadata={"parser": "pymupdf4llm"},
                    )
                )
        return blocks

    @staticmethod
    def _docx_blocks(path: Path) -> List[ParsedTextBlock]:
        from docx import Document

        document = Document(path)
        blocks: List[ParsedTextBlock] = []
        current_heading = "Uploaded document"
        buffer: List[str] = []

        def flush() -> None:
            text = "\n".join(buffer).strip()
            buffer.clear()
            if text:
                blocks.append(
                    ParsedTextBlock(
                        content=text,
                        section=current_heading,
                        metadata={"parser": "python-docx"},
                    )
                )

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            if paragraph.style and paragraph.style.name.startswith("Heading"):
                flush()
                current_heading = text
            else:
                buffer.append(text)
        for table in document.tables:
            rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
            if rows:
                buffer.append("\n".join(rows))
        flush()
        return blocks

    @staticmethod
    def _csv_blocks(path: Path) -> List[ParsedTextBlock]:
        rows: List[str] = []
        with path.open("r", encoding="utf-8-sig", newline="") as handle:
            for index, row in enumerate(csv.reader(handle)):
                if index >= 500:
                    break
                rows.append(" | ".join(cell.strip() for cell in row))
        return [
            ParsedTextBlock(
                content="\n".join(rows),
                section="Uploaded table",
                metadata={"parser": "csv"},
            )
        ] if rows else []

