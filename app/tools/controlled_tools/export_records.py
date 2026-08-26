import hashlib
from html import escape
import inspect
from pathlib import Path
from typing import Any, Dict, List, Literal
from uuid import uuid4

from pydantic import BaseModel, Field

from app.schemas import RiskLevel
from app.tools.definition import (
    ToolCategory,
    ToolDefinition,
    ToolDomain,
    ToolErrorCode,
    ToolExecutionContext,
    ToolPermission,
    ToolResult,
)


class ExportRecordsInput(BaseModel):
    record_ids: List[str] = Field(min_length=1, max_length=50)
    format: Literal["docx", "pdf"] = "docx"
    template_name: Literal["observation", "learning_story", "program_plan", "general"] = "general"


class ExportRecordsOutput(BaseModel):
    export_id: str
    format: str
    template_name: str
    storage_path: str
    checksum: str
    status: str
    expires_at: str | None = None


def build_export_records_tool(store: Any, *, export_root: Path | None = None) -> ToolDefinition:
    root = (export_root or Path("data/exports")).resolve()

    async def async_runtime_handler(input_data: BaseModel, context: ToolExecutionContext) -> ToolResult:
        if not context.teacher_id or not context.class_id:
            return ToolResult.fail(
                code=ToolErrorCode.PERMISSION_DENIED,
                message="Record export requires a trusted teacher and class scope.",
                risk_level=RiskLevel.L3_FORBIDDEN,
                recoverable=False,
            )
        data = ExportRecordsInput.model_validate(input_data)
        records = store.get_exportable_records(
            teacher_id=context.teacher_id,
            class_id=context.class_id,
            record_ids=data.record_ids,
        )
        if inspect.isawaitable(records):
            records = await records
        root.mkdir(parents=True, exist_ok=True)
        target = root / f"{uuid4()}.{data.format}"
        if data.format == "docx":
            _write_docx(target, records, data.template_name)
        else:
            _write_pdf(target, records, data.template_name)
        checksum = hashlib.sha256(target.read_bytes()).hexdigest()
        result = store.save_record_export(
            teacher_id=context.teacher_id,
            class_id=context.class_id,
            record_ids=data.record_ids,
            format=data.format,
            template_name=data.template_name,
            storage_path=str(target),
            checksum=checksum,
        )
        if inspect.isawaitable(result):
            result = await result
        return ToolResult.ok(data=result, risk_level=RiskLevel.L2_CONTROLLED_WRITE)

    return ToolDefinition(
        name="export_records",
        description=(
            "Export already-saved authorised records to DOCX or PDF with a fixed "
            "template. Record IDs must come from an explicit teacher selection, "
            "filter, or trusted workspace reference; never list all records to guess "
            "what an unresolved pronoun means. Never accept arbitrary draft text."
        ),
        category=ToolCategory.DRAFT,
        input_model=ExportRecordsInput,
        output_model=ExportRecordsOutput,
        risk_level=RiskLevel.L2_CONTROLLED_WRITE,
        permission=ToolPermission.REQUIRE_APPROVAL,
        completion_aliases=(
            "导出记录",
            "导出为pdf",
            "导出为docx",
            "export the record",
            "export records",
            "export as pdf",
            "export as docx",
        ),
        domain=ToolDomain.LOCAL,
        parallel_safe=False,
        async_runtime_handler=async_runtime_handler,
    )


def _write_docx(path: Path, records: List[Dict[str, Any]], template_name: str) -> None:
    from docx import Document

    document = Document()
    document.add_heading(f"EasyTeaching — {template_name.replace('_', ' ').title()}", 0)
    for record in records:
        document.add_heading(record.get("title") or "Observation", level=1)
        for label, key in _record_fields(record):
            value = record.get(key)
            if value not in (None, "", []):
                document.add_heading(label, level=2)
                document.add_paragraph(_display(value))
    document.save(path)


def _write_pdf(path: Path, records: List[Dict[str, Any]], template_name: str) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    styles = getSampleStyleSheet()
    story = [Paragraph(f"EasyTeaching — {template_name.replace('_', ' ').title()}", styles["Title"])]
    for record in records:
        story.extend(
            [
                Spacer(1, 12),
                Paragraph(escape(str(record.get("title") or "Observation")), styles["Heading1"]),
            ]
        )
        for label, key in _record_fields(record):
            value = record.get(key)
            if value not in (None, "", []):
                story.append(Paragraph(escape(label), styles["Heading2"]))
                story.append(Paragraph(_display_pdf(value), styles["BodyText"]))
    SimpleDocTemplate(str(path), pagesize=A4).build(story)


def _record_fields(record: Dict[str, Any]):
    if record.get("record_type") == "observation":
        return (
            ("Observed at", "observed_at"),
            ("Setting", "setting"),
            ("Objective observation", "objective_text"),
            ("Educator actions", "educator_actions"),
        )
    return (
        ("Analysis", "analysis"),
        ("Curriculum links", "curriculum_links"),
        ("Next steps", "next_steps"),
    )


def _display(value: Any) -> str:
    if isinstance(value, list):
        return "<br/>".join(f"• {item}" for item in value)
    return str(value).replace("\n", "<br/>")


def _display_pdf(value: Any) -> str:
    if isinstance(value, list):
        return "<br/>".join(f"• {escape(str(item))}" for item in value)
    return escape(str(value)).replace("\n", "<br/>")
